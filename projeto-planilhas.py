from openpyxl import load_workbook
from openpyxl.drawing.image import Image
import random
from docx import Document
from num2words import num2words
import sys

def formatar_reais(valor):
  return f"{valor:,.2f}".replace(",","v").replace(".",",").replace("v",".")
def informacoes_iniciais(escola):
  meses= ("","Janeiro","Fevereiro","Março","Abril","Maio","Junho","Julho","Agosto","Setembro","Outubro","Novembro","Dezembro")
  cnpj_buscar = str(input("Digite o CNPJ da escola que deseja bucar(sem . ou /, ex:123456789) : "))
  nome_escola,cnpj_escola,cep_escola,cidade_escola,endereco_escola,diretor_escola = buscar_escola(escola,cnpj_buscar)
  print(f"\n este cnpj corresponde a escola {nome_escola}\n")
  dia = str(input("Digite a data dos orçamentos (dia): "))
  mes = int(input("Digite a data dos orçamentos (mês): "))
  if mes<1 or mes>12:
    print("o numero do mês está errado")
    sys.exit()
  mes_digitado = meses[mes]
  ano = str(input("Digite a data dos orçamentos (ano): "))
  NF = str(input("Digite o número da nota fiscal correspondente : "))
  return nome_escola,cnpj_escola,cep_escola,cidade_escola,endereco_escola,dia,mes,mes_digitado,ano,NF,diretor_escola,meses

def opcoes():
  dia_emitida = mes_emitida = ano_emitida = None
  dia_consolidacao = mes_consolidacao = ano_consolidacao = None
  dia_recibo = mes_recibo = ano_recibo = None
  dia_emitida = mes_emitida = ano_emitida = None
  meio_pagamento = None
  opcao_consolidacao = str(input("Você quer consolidação? (S/N) "))
  if opcao_consolidacao.lower() == "s":
    dia_consolidacao = int(input("Digite a data da consolidação (dia): "))
    mes_consolidacao = int(input("Digite a data da consolidação (mês): "))
    ano_consolidacao = str(input("Digite a data da consolidação (ano): "))
  elif opcao_consolidacao != "n":
    print("Opção inválida,configurando como NÃO")
    opcao_consolidacao = "n"


  opcao_recibo = str(input("Você quer recibo? (S/N) "))
  if opcao_recibo.lower() == "s":
    dia_recibo = int(input("Digite a data do recibo (dia): "))
    mes_recibo = int(input("Digite a data do recibo (mês): "))
    ano_recibo = str(input("Digite a data do recibo (ano): "))
    nota = str(input("a data que a nota foi emitida é a mesma do recibo? (S/N)"))
    meio_pagamento = str(input("Pago por meio de:  "))
    if nota.lower() == "n":
      dia_emitida = int(input("Digite a data da emissão (dia): "))
      mes_emitida = int(input("Digite a data da emissão (mês): "))
      ano_emitida = int(input("Digite a data da emissão (ano): "))
  elif opcao_recibo != "n":
    print("Opção inválida,configurando como NÃO")
    opcao_recibo = "n"
  return opcao_consolidacao,dia_consolidacao,mes_consolidacao,ano_consolidacao,opcao_recibo,dia_recibo,mes_recibo,ano_recibo,dia_emitida,mes_emitida,ano_emitida,meio_pagamento

def buscar_escola(escola,buscador):
  for row in range(2,escola.max_row + 1):
    celula_buscar = escola.cell(row=row,column = 1).value
    if celula_buscar == buscador:
      nome_escola = escola.cell(row=row,column = 2).value
      cnpj_escola = escola.cell(row=row,column = 3).value
      cep_escola = escola.cell(row=row,column = 4).value
      cidade_escola = escola.cell(row=row,column = 5).value
      endereco_escola = escola.cell(row=row,column = 6).value
      diretor_escola = escola.cell(row=row,column = 7).value
      return nome_escola,cnpj_escola,cep_escola,cidade_escola,endereco_escola,diretor_escola
  print("Escola não encontrada")
  sys.exit()

def aplicar_informacoes(pagina_escolas,nce,paper,grafite):
  nome_escola,cnpj_escola,cep_escola,cidade_escola,endereco_escola,dia,mes,mes_digitado,ano,NF,diretor_escola,meses =informacoes_iniciais(pagina_escolas)
  #nce
  pagina_nce.cell(row=6,column=1).value = f"NOSSA SENHORA DA GLÓRIA/SE {dia} DE {mes_digitado.upper()} DE {ano}"
  pagina_nce.cell(row=8,column=1).value = nome_escola
  pagina_nce.cell(row=9,column=1).value = endereco_escola
  pagina_nce.cell(row=10,column=1).value = f"CEP {cep_escola} - {cidade_escola}"
  pagina_nce.cell(row=11,column=1).value = f"CNPJ: {cnpj_escola}"
  #paper&co
  pagina_paper.cell(row=9,column=1).value = f"NOSSA SENHORA DAS DORES/SE, {dia} DE {mes_digitado.upper()} DE {ano}"
  pagina_paper.cell(row=11,column=1).value = f"{nome_escola}      /      CNPJ- {cnpj_escola}"
  pagina_paper.cell(row=12,column=1).value = endereco_escola
  pagina_paper.cell(row=13,column=1).value = f"CEP {cep_escola}   -   {cidade_escola}"
  #grafite
  pagina_grafite.cell(row=7,column=2).value = nome_escola.title()
  pagina_grafite.cell(row=7,column=5).value = cnpj_escola
  pagina_grafite.cell(row=8,column=2).value = endereco_escola.title()
  pagina_grafite.cell(row=9,column=2).value = cidade_escola.title()
  pagina_grafite.cell(row=10,column=2).value = cep_escola
  pagina_grafite.cell(row=10,column=5).value = f"{dia} DE {mes_digitado.upper()} DE {ano}"

  return dia,mes,ano,NF,diretor_escola,cidade_escola,meses,nome_escola,cnpj_escola

def numero_por_extenso(num):
  try:
    return num2words(num, lang='pt_BR')
  except NotImplementedError:
    return "Número muito grande ou formato não suportado pela biblioteca."
  except Exception as e:
    return f"Ocorreu um erro: {e}"

def fazer_recibo(arquivo_recibo,total_nce,nome_escola,NF,dia_emitida,mes_emitida,ano_emitida,dia_recibo,mes_recibo,ano_recibo,meio_pagamento,meses):
  for paragrafo in arquivo_recibo.paragraphs:
    if "<VALOR>" in paragrafo.text:
      paragrafo.text = paragrafo.text.replace("<VALOR>",str(formatar_reais(total_nce)))
    if "<EXTENSO>" in paragrafo.text:
      paragrafo.text = paragrafo.text.replace("<EXTENSO>",numero_por_extenso(float(total_nce)))
    if "<NOME>" in paragrafo.text:
      paragrafo.text = paragrafo.text.replace("<NOME>",nome_escola)
    if "<NF>" in paragrafo.text:
      paragrafo.text = paragrafo.text.replace("<NF>",str(NF))
    if "<DATA_EMISSÃO>" in paragrafo.text:
      if dia_emitida and mes_emitida and ano_emitida:
        paragrafo.text = paragrafo.text.replace("<DATA_EMISSÃO>",f"{dia_emitida}/{mes_emitida}/{ano_emitida}")
      else:
        paragrafo.text = paragrafo.text.replace("<DATA_EMISSÃO>",f"{dia_recibo}/{mes_recibo}/{ano_recibo}")
    if "<MEIO>" in paragrafo.text:
      paragrafo.text = paragrafo.text.replace("<MEIO>",meio_pagamento)
    if "<DATA>" in paragrafo.text:
      paragrafo.text = paragrafo.text.replace("<DATA>",f"{dia_recibo}/{mes_recibo}/{ano_recibo}")
    if "<DIA>" in paragrafo.text:
      paragrafo.text = paragrafo.text.replace("<DIA>",str(dia_recibo))
    if "<MÊS>" in paragrafo.text:
      paragrafo.text = paragrafo.text.replace("<MÊS>",str(meses[mes_recibo]))
    if "<ANO>" in paragrafo.text:
      paragrafo.text = paragrafo.text.replace("<ANO>",str(ano_recibo))

def fazer_consolidacao(arquivo_consolidacao,item_numero,produto,un,qt,unit_nce,unit_paper,unit_grafite,diretor_escola,
                       nome_escola,dia_consolidacao,mes_consolidacao,ano_consolidacao,cidade_escola,cnpj_escola,total_nce,total_paper,total_grafite,meses):
    for paragrafo in arquivo_consolidacao.paragraphs:
      if f"<VALOR{item_numero}>" in paragrafo.text:
        paragrafo.text = paragrafo.text.replace(f"<VALOR{item_numero}",str(produto))
      if f"<UNI{item_numero}>" in paragrafo.text:
        paragrafo.text = paragrafo.text.replace(f"<UNI{item_numero}",str(un))
      if f"<Q{item_numero}>" in paragrafo.text:
        paragrafo.text = paragrafo.text.replace(f"<Q{item_numero}",str(qt))
      if f"<VALOR_A{item_numero}>" in paragrafo.text:
        paragrafo.text = paragrafo.text.replace(f"<VALOR_A{item_numero}",str(unit_nce))
      if f"<VALOR_B{item_numero}>" in paragrafo.text:
        paragrafo.text = paragrafo.text.replace(f"<VALOR_B{item_numero}",str(unit_paper))
      if f"<VALOR_C{item_numero}>" in paragrafo.text:
        paragrafo.text = paragrafo.text.replace(f"<VALOR_C{item_numero}",str(unit_grafite))
      if "<DIRETOR>" in paragrafo.text:
        paragrafo.text = paragrafo.text.replace("DIRETOR",diretor_escola)
      if "<CIDADE>" in paragrafo.text:
        paragrafo.text = paragrafo.text.replace("CIDADE",cidade_escola)
      if "<DATA>" in paragrafo.text:
        paragrafo.text = paragrafo.text.replace("DATA",f"{dia_consolidacao} de {meses[mes_consolidacao]} de {ano_consolidacao}")
      if "<TOTAL_A>" in paragrafo.text:
        paragrafo.text = paragrafo.text.replace("TOTAL_A",str(formatar_reais(total_nce)))
      if "<TOTAL_B>" in paragrafo.text:
        paragrafo.text = paragrafo.text.replace("TOTAL_B",str(formatar_reais(total_paper)))
      if "<TOTAL_C>" in paragrafo.text:
        paragrafo.text = paragrafo.text.replace("TOTAL_C",str(formatar_reais(total_grafite)))
      if "<NOME>" in paragrafo.text:
        paragrafo.text = paragrafo.text.replace("NOME",nome_escola)
      if "<CNPJ>" in paragrafo.text:
        paragrafo.text = paragrafo.text.replace("CNPJ",cnpj_escola)


#abertura de arquivos
try:
  arquivo_nce = load_workbook("MODELO NCE.xlsx")
except FileNotFoundError:
  print("Arquivo MODELO NCE não encontrado. Verifique o nome e tente novamente.")
  sys.exit()
pagina_nce = arquivo_nce["Matriz"]
img_nce1 = Image("logos/nce1.png")
img_nce2 = Image("logos/nce2.jpeg")
img_nce1.height = 56
img_nce1.width = 320
img_nce2.height = 51
img_nce2.width = 96
pagina_nce.add_image(img_nce1,"D1")
pagina_nce.add_image(img_nce2,"B2")


try:
  arquivo_paper = load_workbook("MODELO PAPER.xlsx")
except FileNotFoundError:
  print(f"Arquivo MODELO PAPER não encontrado. Verifique o nome e tente novamente.")
  sys.exit()
pagina_paper = arquivo_paper.active
img_paper = Image("logos/paper.png")
img_paper.width = 196
img_paper.height = 80
pagina_paper.add_image(img_paper,"B1")

try:
  arquivo_grafite = load_workbook("MODELO GRAFITE.xlsx")
except FileNotFoundError:
  print(f"Arquivo MODELO GRAFITE não encontrado. Verifique o nome e tente novamente.")
  sys.exit()
pagina_grafite = arquivo_grafite.active
img_grafite = Image("logos/grafite.jpeg")
img_grafite.width = 608
img_grafite.height = 48
pagina_grafite.add_image(img_grafite,"A1")
try:
  arquivo_consolidacao = Document("MODELO CONSOLIDACAO.docx")
except FileNotFoundError:
  print(f"Arquivo MODELO CONSOLIDACAO não encontrado. Verifique o nome e tente novamente.")
  sys.exit()

try:
  arquivo_recibo = Document("MODELO RECIBO.docx")
except FileNotFoundError:
  print(f"Arquivo MODELO RECIBO não encontrado. Verifique o nome e tente novamente.")
  sys.exit()

try:
  arquivo_controle = load_workbook("MODELO CONTROLE.xlsx")
except FileNotFoundError:
  print(f"Arquivo MODELO CONTROLE não encontrado. Verifique o nome e tente novamente.")
  sys.exit()
pagina_controle = arquivo_controle.active

try:
  arquivo_escolas = load_workbook("escolas.xlsx")
except FileNotFoundError:
  print(f"Arquivo escolas.xlsx não encontrado. Verifique o nome e tente novamente.")
  sys.exit()
pagina_escolas = arquivo_escolas.active
dia,mes,ano,NF,diretor_escola,cidade_escola,meses,nome_escola,cnpj_escola = aplicar_informacoes(pagina_escolas,pagina_nce,pagina_paper,pagina_grafite)
opcao_consolidacao,dia_consolidacao,mes_consolidacao,ano_consolidacao,opcao_recibo,dia_recibo,mes_recibo,ano_recibo,dia_emitida,mes_emitida,ano_emitida,meio_pagamento = opcoes()

nome_doador = input("Digite o nome do arquivo do doador (sem .xlsx): ")
try:
  arquivo_doador = load_workbook(f"doadores/{nome_doador}.xlsx")
except FileNotFoundError:
  print(f"Arquivo '{nome_doador}' não encontrado. Verifique o nome e tente novamente.")
  sys.exit()
pagina_doador = arquivo_doador["Table 10"]




linha_ini_nce = 15
linha_ini_paper = 17
linha_ini_grafite = 12
linha_ini_controle = 3
linha_doadora = 3

precos_randomizados_por_valor = {}

total_nce = 0
total_paper = 0
total_grafite = 0
item_numero = 1
while pagina_doador.cell(row=linha_doadora,column=2).value is not None:
  print(f"Processando produto {item_numero} . . .\n")
  produto = pagina_doador.cell(row=linha_doadora,column=2).value
  un = pagina_doador.cell(row=linha_doadora,column=6).value
  qt = int(pagina_doador.cell(row=linha_doadora,column=7).value)
  unit_nce = float(pagina_doador.cell(row=linha_doadora,column=8).value)

  if unit_nce in precos_randomizados_por_valor:
    unit_paper,unit_grafite = precos_randomizados_por_valor[unit_nce]
  else:
   unit_paper = round(unit_nce * (1+ random.uniform(0.15,0.25))*20)/20
   unit_grafite = round(unit_nce * (1+ random.uniform(0.15,0.25))*20)/20
   while abs(unit_paper - unit_grafite)< 0.04:
     unit_grafite = unit_nce * (1+ random.uniform(0.15,0.25))
   precos_randomizados_por_valor[unit_nce] = unit_paper,unit_grafite


  pagina_nce.cell(row=linha_ini_nce,column=2).value = produto
  pagina_nce.cell(row=linha_ini_nce,column=3).value = un
  pagina_nce.cell(row=linha_ini_nce,column=4).value = qt
  pagina_nce.cell(row=linha_ini_nce,column=5).value = unit_nce
  total_nce = total_nce +(unit_nce*qt)


  pagina_paper.cell(row=linha_ini_paper,column=2).value = produto
  pagina_paper.cell(row=linha_ini_paper,column=3).value = un
  pagina_paper.cell(row=linha_ini_paper,column=4).value = qt
  pagina_paper.cell(row=linha_ini_paper,column=5).value = unit_paper
  total_paper = total_paper +(unit_paper*qt)

  pagina_grafite.cell(row=linha_ini_grafite,column=2).value = produto
  pagina_grafite.cell(row=linha_ini_grafite,column=3).value = un
  pagina_grafite.cell(row=linha_ini_grafite,column=4).value = qt
  pagina_grafite.cell(row=linha_ini_grafite,column=5).value = unit_grafite
  total_grafite = total_grafite +(unit_grafite*qt)

  pagina_controle.cell(row=linha_ini_controle,column=2).value = produto
  pagina_controle.cell(row=linha_ini_controle,column=3).value = un
  pagina_controle.cell(row=linha_ini_controle,column=4).value = qt
  pagina_controle.cell(row=linha_ini_controle,column=5).value = unit_nce
  pagina_controle.cell(row=linha_ini_controle,column=8).value = unit_paper
  pagina_controle.cell(row=linha_ini_controle,column=11).value = unit_grafite

  #CONSOLIDAÇÃO
  if opcao_consolidacao.lower() == "s":
     fazer_consolidacao(arquivo_consolidacao,item_numero,produto,un,qt,unit_nce,unit_paper,unit_grafite,diretor_escola,
                        nome_escola,dia_consolidacao,mes_consolidacao,ano_consolidacao,cidade_escola,cnpj_escola,total_nce,total_paper,total_grafite,meses)

  item_numero += 1
  linha_doadora +=1
  linha_ini_grafite +=1
  linha_ini_paper +=1
  linha_ini_nce +=1
  linha_ini_controle +=1

if opcao_recibo.lower() == "s":
  fazer_recibo(arquivo_recibo,total_nce,nome_escola,NF,dia_emitida,mes_emitida,ano_emitida,dia_recibo,mes_recibo,ano_recibo,meio_pagamento,meses)



print("\nProcesso finalizado!")

arquivo_nce.save(f"arquivos/ORÇAMENTO NF{NF} {ano}-{mes}-{dia} NCE.xlsx")
arquivo_paper.save(f"arquivos/ORÇAMENTO NF{NF} {ano}-{mes}-{dia} PAPER&CO.xlsx")
arquivo_grafite.save(f"arquivos/ORÇAMENTO NF{NF} {ano}-{mes}-{dia} GRAFITE.xlsx")
arquivo_controle.save(f"arquivos/MODELO DOC NF{NF}.xlsx")
if opcao_consolidacao.lower() == "s":
  arquivo_consolidacao.save(f"arquivos/ORÇAMENTO NF{NF} {ano_consolidacao}-{mes_consolidacao}-{dia_consolidacao} CONSOLIDAÇÃO DE PESQ DE PREÇO")
if opcao_recibo.lower() == "s":
  arquivo_recibo.save(f"arquivos/ORÇAMENTO NF{NF} {ano_recibo}-{mes_recibo}-{dia_recibo} RECIBO {round(total_nce,2)} NCE")
